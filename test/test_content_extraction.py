#!/usr/bin/env python3
"""
测试不同文件格式的内容提取功能
"""

from utils.logger import logger
from database.models import DatabaseManager
from api.upload import extract_file_content
import asyncio
import os
import sys
import tempfile
from pathlib import Path

# 添加项目根目录到Python路径
sys.path.append(str(Path(__file__).parent.parent))


try:
    from docx import Document
    from docx.shared import Inches
    docx_available = True
except ImportError:
    docx_available = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    pdf_available = True
except ImportError:
    pdf_available = False


async def test_file_content_extraction():
    """测试不同文件格式的内容提取"""
    logger.info("🧪 开始测试文件内容提取功能...")

    # 创建测试目录
    test_dir = Path("test/temp_content_test")
    test_dir.mkdir(exist_ok=True)

    try:
        # 1. 测试Markdown文件
        logger.info("📝 测试Markdown文件...")
        md_file = test_dir / "test.md"
        md_content = """# 测试文档

这是一个测试Markdown文档。

## 功能列表
- 文本提取
- 向量化
- 搜索

**粗体文本** 和 *斜体文本*"""

        md_file.write_text(md_content, encoding='utf-8')
        extracted_md = await extract_file_content(md_file, ".md")

        if extracted_md.strip() == md_content.strip():
            logger.info("✅ Markdown文件内容提取成功")
        else:
            logger.error("❌ Markdown文件内容提取失败")
            logger.error(f"期望内容: {repr(md_content)}")
            logger.error(f"提取内容: {repr(extracted_md)}")

        # 2. 测试DOCX文件（如果库可用）
        if docx_available:
            logger.info("📄 测试DOCX文件...")
            docx_file = test_dir / "test.docx"
            docx_content = "这是一个测试DOCX文档。\n包含多个段落。\n支持向量化功能。"

            # 创建DOCX文件
            doc = Document()
            doc.add_heading('测试文档', 0)
            doc.add_paragraph('这是一个测试DOCX文档。')
            doc.add_paragraph('包含多个段落。')
            doc.add_paragraph('支持向量化功能。')
            doc.save(docx_file)

            extracted_docx = await extract_file_content(docx_file, ".docx")

            if "测试DOCX文档" in extracted_docx and "向量化功能" in extracted_docx:
                logger.info("✅ DOCX文件内容提取成功")
                logger.info(f"提取的内容: {extracted_docx[:100]}...")
            else:
                logger.error("❌ DOCX文件内容提取失败")
                logger.error(f"提取内容: {repr(extracted_docx)}")
        else:
            logger.warning("⚠️ python-docx库未安装，跳过DOCX测试")

        # 3. 测试PDF文件（如果库可用）
        if pdf_available:
            logger.info("📑 测试PDF文件...")
            pdf_file = test_dir / "test.pdf"

            # 创建PDF文件
            c = canvas.Canvas(str(pdf_file), pagesize=letter)
            width, height = letter

            c.drawString(72, height - 72, "测试PDF文档")
            c.drawString(72, height - 100, "这是一个测试PDF文档。")
            c.drawString(72, height - 128, "包含中文内容。")
            c.drawString(72, height - 156, "支持向量化功能。")
            c.save()

            extracted_pdf = await extract_file_content(pdf_file, ".pdf")

            if "测试PDF文档" in extracted_pdf and "向量化功能" in extracted_pdf:
                logger.info("✅ PDF文件内容提取成功")
                logger.info(f"提取的内容: {extracted_pdf[:100]}...")
            else:
                logger.error("❌ PDF文件内容提取失败")
                logger.error(f"提取内容: {repr(extracted_pdf)}")
        else:
            logger.warning("⚠️ reportlab库未安装，无法创建测试PDF文件")
            logger.info("📑 测试现有PDF文件解析...")

            # 创建一个简单的文本作为PDF内容测试
            simple_pdf_content = "测试PDF内容"
            extracted = await extract_file_content(Path("nonexistent.pdf"), ".pdf")
            if "PDF处理库未安装" in extracted or "提取文件内容失败" in extracted:
                logger.info("✅ PDF错误处理正常")
            else:
                logger.error(f"❌ PDF错误处理异常: {extracted}")

        # 4. 测试不支持的文件类型
        logger.info("❓ 测试不支持的文件类型...")
        unsupported_result = await extract_file_content(Path("test.xyz"), ".xyz")
        if "不支持的文件类型" in unsupported_result:
            logger.info("✅ 不支持文件类型的错误处理正常")
        else:
            logger.error(f"❌ 不支持文件类型的错误处理异常: {unsupported_result}")

    except Exception as e:
        logger.error(f"❌ 测试过程中发生错误: {e}")
        raise
    finally:
        # 清理测试文件
        import shutil
        if test_dir.exists():
            shutil.rmtree(test_dir)
            logger.info("🗑️ 清理测试文件")


async def main():
    """主函数"""
    try:
        await test_file_content_extraction()
        logger.info("✅ 文件内容提取测试完成")
    except Exception as e:
        logger.error(f"❌ 测试失败: {e}")
        raise


if __name__ == "__main__":
    asyncio.run(main())
