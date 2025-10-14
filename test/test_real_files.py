#!/usr/bin/env python3
"""
测试实际文件的内容提取功能
"""

from utils.logger import logger
from database.models import DatabaseManager
from api.upload import extract_file_content, validate_file_type
import asyncio
import sys
from pathlib import Path

# 添加项目根目录到Python路径
sys.path.append(str(Path(__file__).parent.parent))


try:
    from docx import Document
    docx_available = True
except ImportError:
    docx_available = False

try:
    import PyPDF2
    pdf_available = True
except ImportError:
    pdf_available = False


async def create_test_files():
    """创建各种格式的测试文件"""
    test_dir = Path("test/temp_real_files")
    test_dir.mkdir(exist_ok=True)

    # 1. 创建Markdown文件
    md_file = test_dir / "sample.md"
    md_content = """# RAG系统文档

## 简介
这是一个检索增强生成(RAG)系统，用于处理文档向量化和智能问答。

## 主要功能
1. **文档上传**: 支持PDF、DOCX、Markdown格式
2. **向量化处理**: 使用OpenAI API进行文本向量化
3. **智能检索**: 基于向量相似度的文档检索
4. **任务管理**: 异步处理向量化任务队列

## 技术架构
- **后端**: FastAPI + SQLite + aiosqlite
- **向量化**: OpenAI Embeddings API
- **文档处理**: python-docx, PyPDF2
- **任务队列**: asyncio + 线程池

## 使用方法
```python
# 上传文件
POST /api/v1/files

# 查看文件内容
GET /api/v1/files/{file_id}

# 删除文件
DELETE /api/v1/files/{file_id}
```

> 注意: 系统支持中文文档处理
"""

    md_file.write_text(md_content, encoding='utf-8')
    logger.info(f"✅ 创建Markdown文件: {md_file}")

    # 2. 创建DOCX文件
    if docx_available:
        docx_file = test_dir / "sample.docx"
        doc = Document()

        # 添加标题
        doc.add_heading('RAG系统用户手册', 0)

        # 添加段落
        doc.add_heading('系统概述', level=1)
        doc.add_paragraph(
            'RAG (Retrieval-Augmented Generation) 系统是一个基于FastAPI开发的文档处理和向量化平台。')

        doc.add_heading('主要特性', level=1)
        features = doc.add_paragraph()
        features.add_run('• 支持多种文档格式\n')
        features.add_run('• 异步向量化处理\n')
        features.add_run('• 智能文档检索\n')
        features.add_run('• 完整的API接口')

        doc.add_heading('安装配置', level=1)
        doc.add_paragraph('使用Poetry进行依赖管理：')
        doc.add_paragraph('poetry install', style='Intense Quote')

        doc.add_heading('API使用', level=1)
        doc.add_paragraph(
            '系统提供RESTful API接口，支持文件上传、内容查看和删除操作。所有接口都返回统一的JSON响应格式。')

        doc.save(docx_file)
        logger.info(f"✅ 创建DOCX文件: {docx_file}")

    return test_dir


async def test_real_file_extraction():
    """测试真实文件的内容提取"""
    logger.info("🧪 开始测试真实文件内容提取...")

    # 创建测试文件
    test_dir = await create_test_files()

    try:
        # 测试所有文件
        test_files = list(test_dir.glob("*"))

        for file_path in test_files:
            if file_path.is_file():
                file_ext = file_path.suffix.lower()
                logger.info(f"\n📄 测试文件: {file_path.name} ({file_ext})")

                # 验证文件类型
                with open(file_path, 'rb') as f:
                    file_content = f.read()

                if validate_file_type(file_path.name, file_content):
                    logger.info("✅ 文件类型验证通过")

                    # 提取内容
                    extracted_content = await extract_file_content(file_path, file_ext)

                    if extracted_content and "提取文件内容失败" not in extracted_content:
                        logger.info("✅ 内容提取成功")
                        logger.info(f"📝 内容长度: {len(extracted_content)} 字符")

                        # 显示前200字符
                        preview = extracted_content[:200].replace('\n', ' ')
                        logger.info(f"📄 内容预览: {preview}...")

                        # 检查关键词
                        keywords = ["RAG", "系统", "向量化", "API", "文档"]
                        found_keywords = [
                            kw for kw in keywords if kw in extracted_content]
                        logger.info(f"🔍 找到关键词: {found_keywords}")

                    else:
                        logger.error(f"❌ 内容提取失败: {extracted_content}")
                else:
                    logger.error("❌ 文件类型验证失败")

        # 测试数据库集成
        logger.info(f"\n🗄️ 测试数据库集成...")

        db_manager = DatabaseManager()
        await db_manager.init_database()

        # 模拟文件上传到数据库
        md_file = test_dir / "sample.md"
        if md_file.exists():
            await db_manager.insert_file(
                file_id="test_md_content",
                original_name="sample.md",
                file_name="sample.md",
                file_path=str(md_file),
                file_type=".md",
                file_size=md_file.stat().st_size
            )

            logger.info(f"✅ 文件记录创建成功: test_md_content")

            # 通过API获取文件内容
            file_info = await db_manager.get_file_by_id("test_md_content")
            if file_info:
                content = await extract_file_content(Path(file_info["file_path"]), file_info["file_type"])
                if content and len(content) > 100:
                    logger.info("✅ 数据库集成测试成功")
                else:
                    logger.error("❌ 数据库集成测试失败")

            # 清理测试记录
            await db_manager.delete_file_and_documents("test_md_content")
            logger.info("🗑️ 清理测试记录")

    except Exception as e:
        logger.error(f"❌ 测试过程中发生错误: {e}")
        raise
    finally:
        # 清理测试文件
        import shutil
        if test_dir.exists():
            shutil.rmtree(test_dir)
            logger.info("🗑️ 清理测试文件")


async def main():
    """主函数"""
    try:
        await test_real_file_extraction()
        logger.info("✅ 真实文件内容提取测试完成")
    except Exception as e:
        logger.error(f"❌ 测试失败: {e}")
        raise


if __name__ == "__main__":
    asyncio.run(main())
